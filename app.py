import streamlit as st
import pandas as pd
import os
import re
import zipfile
import io
import time
import requests
import json
import numpy as np
import docx 
import pickle
import hashlib
import xml.etree.ElementTree as ET
from urllib.parse import urljoin, quote, unquote
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from concurrent.futures import ThreadPoolExecutor, as_completed
from tenacity import retry, stop_after_attempt, wait_exponential

# ====================
# Configuration & Constants
# ====================
CACHE_DIR = "vector_store_cache"
if not os.path.exists(CACHE_DIR):
    os.makedirs(CACHE_DIR)

# ====================
# Lightweight WebDAV Client (With Folder Support)
# ====================

class SimpleWebDavClient:
    """基于 requests 的轻量级 WebDAV 客户端"""
    def __init__(self, base_url, username, password):
        self.base_url = base_url.rstrip('/') 
        self.auth = (username, password)
        self.session = requests.Session()
        self.session.auth = self.auth
    
    def list(self, path="/"):
        """列出指定路径下的文件和文件夹"""
        path = path.strip('/')
        full_url = f"{self.base_url}/{path}/" if path else f"{self.base_url}/"
        
        headers = {'Depth': '1'}
        try:
            response = self.session.request('PROPFIND', full_url, headers=headers)
            if response.status_code in [200, 207]:
                return self._parse_propfind(response.content, full_url)
            else:
                raise Exception(f"WebDAV Error: {response.status_code} - {response.text}")
        except Exception as e:
            raise Exception(f"Connection failed: {str(e)}")

    def download(self, path):
        """下载文件内容"""
        # path 是相对于 base_url 的路径
        full_url = f"{self.base_url}/{quote(path.strip('/'))}"
        response = self.session.get(full_url)
        if response.status_code == 200:
            return response.content
        else:
            raise Exception(f"Download failed: {response.status_code}")

    def _parse_propfind(self, xml_content, current_url):
        """解析 XML 响应获取文件列表，区分文件和文件夹"""
        items = []
        try:
            root = ET.fromstring(xml_content)
            # 处理 namespace，WebDAV 响应通常带有复杂的 namespace
            # 简单处理：忽略 namespace 直接查找 local name
            for response in root.findall('.//{DAV:}response'):
                href = response.find('.//{DAV:}href').text
                href = unquote(href)
                
                # 判断是否是集合 (文件夹)
                resourcetype = response.find('.//{DAV:}resourcetype')
                is_collection = False
                if resourcetype is not None:
                    if resourcetype.find('.//{DAV:}collection') is not None:
                        is_collection = True
                
                # 提取相对路径/名称
                # href 通常是 /dav/folder/file.docx
                # 我们需要提取出显示名称
                name = href.rstrip('/').split('/')[-1]
                
                # 过滤掉当前目录本身
                # 比较 href 和 current_url 的路径部分是否一致
                # 这里简单比对 name 是否为空 (根目录) 或 href 是否等于 current_url
                
                if not name: continue 
                
                items.append({
                    'name': name,
                    'path': href, # 保留完整 href 用于导航
                    'is_folder': is_collection
                })
        except Exception as e:
            print(f"XML Parsing Error: {e}")
            pass
        
        # 过滤掉当前目录自己 (通常它是列表的第一个，且名字和当前目录一样)
        # 这里做一个简单的去重逻辑：如果列表里有 path 结尾和请求 path 一样的，去掉
        return items

# ====================
# Core Classes for AI & Data
# ====================

class LLMClient:
    """处理与大模型的交互 (Embedding 和 Chat)"""
    def __init__(self, config):
        self.base_url = config.get('base_url', '').rstrip('/')
        self.api_key = config.get('api_key')
        self.model = config.get('model')
        self.embedding_model = config.get('embedding_model', 'text-embedding-004')
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.embedding_failed = False 

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    def get_embedding(self, text):
        """获取文本的向量表示"""
        if self.embedding_failed: return None
        
        payload = {
            "input": text.replace("\n", " "),
            "model": self.embedding_model
        }
        url = f"{self.base_url}/embeddings"
        
        try:
            response = requests.post(url, headers=self.headers, json=payload, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if 'data' in data and len(data['data']) > 0:
                    return data['data'][0]['embedding']
            else:
                print(f"Embedding failed: {response.status_code}")
                return None
        except Exception as e:
            print(f"Embedding error: {e}")
            return None

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    def chat_completion(self, system_prompt, user_prompt):
        """调用 Chat 接口"""
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.1,
            "response_format": {"type": "json_object"}
        }
        try:
            response = requests.post(f"{self.base_url}/chat/completions", headers=self.headers, json=payload, timeout=60)
            if response.status_code == 200:
                content = response.json()['choices'][0]['message']['content']
                content = content.replace("```json", "").replace("```", "")
                return json.loads(content)
            else:
                raise Exception(f"API Error {response.status_code}: {response.text}")
        except Exception as e:
            raise Exception(f"Request failed: {str(e)}")
            
    def test_connection(self):
        results = {"chat": False, "embedding": False, "msg": ""}
        try:
            self.chat_completion("Test", "Reply JSON: {'status': 'ok'}")
            results['chat'] = True
        except Exception as e: results['msg'] += f"Chat Error: {e}\n"
        try:
            if self.get_embedding("test"): results['embedding'] = True
            else: results['msg'] += "Embedding Error: None (Check model)\n"
        except Exception as e: results['msg'] += f"Embedding Exception: {e}\n"
        return results

class VectorStore:
    """持久化向量数据库 (支持制度库 + 法规库)"""
    def __init__(self):
        self.documents = []   # 制度文档
        self.vectors = []     # 制度向量
        self.regulations = [] # 法规文档 (缓存用)
        self.llm_client = None

    def set_client(self, client):
        self.llm_client = client

    def save_to_disk(self, name="default"):
        """保存完整库 (制度 + 法规)"""
        path = os.path.join(CACHE_DIR, f"{name}.pkl")
        data = {
            "documents": self.documents,
            "vectors": self.vectors,
            "regulations": self.regulations
        }
        with open(path, "wb") as f:
            pickle.dump(data, f)
        return path

    def load_from_disk(self, name="default"):
        """加载库"""
        path = os.path.join(CACHE_DIR, f"{name}.pkl")
        if os.path.exists(path):
            with open(path, "rb") as f:
                data = pickle.load(f)
                self.documents = data.get("documents", [])
                self.vectors = data.get("vectors", [])
                self.regulations = data.get("regulations", [])
            return True
        return False

    def add_documents(self, file_corpus):
        """添加制度文档并向量化"""
        chunk_size = 500
        overlap = 50 
        new_docs = []
        texts_to_embed = []
        start_doc_id = len(self.documents)
        
        for file in file_corpus:
            if any(d['source'] == file['name'] for d in self.documents): continue
            
            text = file['content']
            for i in range(0, len(text), chunk_size - overlap):
                chunk = text[i:i + chunk_size]
                if len(chunk) < 50: continue 
                keywords = set(re.split(r'[，。；：\s]', chunk))
                keywords = [k for k in keywords if len(k) > 1]
                new_docs.append({'id': start_doc_id, 'text': chunk, 'source': file['name'], 'keywords': keywords})
                texts_to_embed.append(chunk)
                start_doc_id += 1
        
        if not texts_to_embed: return

        if self.llm_client:
            with st.status(f"正在向量化 {len(texts_to_embed)} 个新制度片段...") as status:
                new_vectors = [None] * len(texts_to_embed)
                with ThreadPoolExecutor(max_workers=5) as executor:
                    future_to_idx = {executor.submit(self.llm_client.get_embedding, t): i for i, t in enumerate(texts_to_embed)}
                    for future in as_completed(future_to_idx):
                        idx = future_to_idx[future]
                        new_vectors[idx] = future.result()
                
                self.documents.extend(new_docs)
                self.vectors = (list(self.vectors) if len(self.vectors)>0 else []) + new_vectors
                status.update(label="制度入库完成", state="complete")

    def add_regulations(self, file_corpus):
        """添加法规文档 (无需向量化，仅解析保存)"""
        count = 0
        for file in file_corpus:
            # 查重
            if any(r['name'] == file['name'] for r in self.regulations): continue
            self.regulations.append(file) # file: {'name': name, 'content': text}
            count += 1
        return count

    def search(self, query_text, top_k=3):
        """混合检索 (仅针对制度库)"""
        vec_results = []
        valid_indices = [i for i, v in enumerate(self.vectors) if v is not None]
        
        if valid_indices and self.llm_client:
            query_vec = self.llm_client.get_embedding(query_text)
            if query_vec is not None:
                q_v = np.array(query_vec)
                norm_q = np.linalg.norm(q_v)
                matrix = np.array([self.vectors[i] for i in valid_indices])
                norm_matrix = np.linalg.norm(matrix, axis=1)
                if norm_q > 0:
                    scores = np.dot(matrix, q_v) / (norm_matrix * norm_q)
                    top_k_indices = np.argsort(scores)[-top_k:][::-1]
                    for idx_in_valid in top_k_indices:
                        real_idx = valid_indices[idx_in_valid]
                        score = scores[idx_in_valid]
                        if score > 0: vec_results.append({'doc': self.documents[real_idx], 'score': float(score), 'method': 'vector'})

        kw_results = []
        query_keywords = [k for k in re.split(r'[，。；：\s]', query_text) if len(k) > 1]
        for doc in self.documents:
            overlap = sum(1 for k in query_keywords if k in doc['keywords']) 
            if overlap > 0:
                score = overlap / (len(query_keywords) + 1) * 0.8 
                kw_results.append({'doc': doc, 'score': score, 'method': 'keyword'})
        
        kw_results.sort(key=lambda x: x['score'], reverse=True)
        kw_results = kw_results[:top_k]

        combined = vec_results + kw_results
        seen_ids = set()
        final_results = []
        combined.sort(key=lambda x: x['score'], reverse=True)
        for res in combined:
            did = res['doc']['id']
            if did not in seen_ids:
                final_results.append({'source': res['doc']['source'], 'content': res['doc']['text'], 'score': res['score']})
                seen_ids.add(did)
            if len(final_results) >= top_k: break
        return final_results

# ====================
# Helper Functions
# ====================

def process_uploaded_files(uploaded_files):
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith('.zip'):
            try:
                with zipfile.ZipFile(uploaded_file) as z:
                    for filename in z.namelist():
                        if filename.endswith('/') or filename.startswith('__MACOSX') or filename.startswith('._'): continue
                        if filename.endswith(('.docx', '.xlsx')):
                            with z.open(filename) as f: yield filename, f.read()
            except Exception as e: st.error(f"解压失败: {e}")
        else: yield uploaded_file.name, uploaded_file.getvalue()

def extract_text_from_content(filename, content):
    text = ""
    try:
        file_stream = io.BytesIO(content)
        if filename.endswith('.docx'):
            doc = docx.Document(file_stream)
            text = '\n'.join([para.text for para in doc.paragraphs])
        elif filename.endswith('.xlsx'):
            df_dict = pd.read_excel(file_stream, sheet_name=None, header=None)
            text_parts = []
            for sheet_name, df in df_dict.items():
                sheet_text = df.astype(str).apply(lambda x: ' '.join(x), axis=1)
                text_parts.append('\n'.join(sheet_text))
            text = '\n'.join(text_parts)
    except Exception as e: print(f"Error parsing {filename}: {e}")
    return text

def parse_regulation_clauses(text):
    pattern = r'(第\s*[\d零一二三四五六七八九十百]+\s*条|Article\s+\d+)'
    parts = re.split(pattern, text)
    clauses = []
    if len(parts) > 1:
        for i in range(1, len(parts), 2):
            title = parts[i].strip()
            content = parts[i+1].strip() if i+1 < len(parts) else ""
            clauses.append({"条款号": title, "法规正文": title + " " + content, "适用性": "适用"})
    return clauses

def evaluate_single_clause(clause, vector_store, llm_client):
    row = {"条款号": clause['条款号'], "法规正文": clause['法规正文'], "评价结论": "❌缺失/不符合", "差距分析": "未检索到相关制度", "改进建议": "请补充相关管理规定", "支撑证据": "无", "匹配度": 0.0}

    search_results = vector_store.search(clause['法规正文'], top_k=3)
    top_score = search_results[0]['score'] if search_results else 0
    row['匹配度'] = top_score
        
    evidence_text = ""
    if search_results:
        for i, res in enumerate(search_results):
            evidence_text += f"参考制度片段 {i+1} (来源: {res['source']}):\n{res['content'][:800]}\n---\n"
    else:
        evidence_text = "未检索到任何相关的企业内部制度文档。"
    
    system_prompt = """你是一名具有20年经验的EHS管理专家，精通中国EHS法规标准，擅长仓储物流场景。
    请对给定的法规条款进行合规性评价。严格执行以下思维链：
    1. 解读：理解条款核心要求（人机料法环），判定是否适用于物流仓储企业。如果不适用，直接标记“不适用”。
    2. 比对：对比法规要求与提供的企业制度片段。是否覆盖所有要素？针对物流场景是否具体可执行？
    3. 判定：给出定性结论。"""
    
    user_prompt = f"""
    【法规条款】
    {clause['法规正文']}

    【企业制度现状（检索到的最相关片段）】
    {evidence_text}

    【任务要求】
    请严格按照以下格式返回 JSON 结果：
    {{
        "applicability": "适用" 或 "不适用",
        "compliance_status": "完全符合" 或 "部分符合/需完善" 或 "缺失/不符合" 或 "不适用",
        "gap_analysis": "150字以内。若不适用填'不适用'。若适用，分析具体缺了什么或说明合规点。",
        "improvement_suggestion": "若完全符合填'无'。否则给出1-2条建议。",
        "evidence_summary": "列出最匹配的制度名称及关键句摘要"
    }}
    """
    try:
        result = llm_client.chat_completion(system_prompt, user_prompt)
        status = result.get('compliance_status', '缺失/不符合')
        if "完全符合" in status: status = "✅完全符合"
        elif "部分" in status or "需完善" in status: status = "⚠️部分符合/需完善"
        elif "不适用" in status: status = "❗不适用"
        else: status = "❌缺失/不符合"
        
        row['评价结论'] = status
        row['差距分析'] = result.get('gap_analysis', '无分析')
        row['改进建议'] = result.get('improvement_suggestion', '无建议')
        row['支撑证据'] = result.get('evidence_summary', '无')
    except Exception as e:
        row['差距分析'] = f"LLM分析失败: {str(e)}"
    return row

def generate_word_report(df_results, summary_stats):
    doc = Document()
    title = doc.add_heading('EHS法规合规性评价报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"评价日期: {time.strftime('%Y-%m-%d')}")
    
    # 1. 总体评价 - 系统性汇总
    doc.add_heading('第一部分：总体评价与管理建议', level=1)
    
    # 1.1 数据概览
    doc.add_heading('1.1 评价数据概览', level=2)
    p = doc.add_paragraph()
    p.add_run(f"本次评价共分析法规条款 {summary_stats['total']} 条，其中：\n")
    p.add_run(f"✅ 完全符合: {summary_stats['compliant']} 条 ({summary_stats['compliant']/summary_stats['total']*100:.1f}%)\n").font.color.rgb = RGBColor(0, 128, 0)
    p.add_run(f"⚠️ 部分符合/需完善: {summary_stats['partial']} 条 ({summary_stats['partial']/summary_stats['total']*100:.1f}%)\n").font.color.rgb = RGBColor(255, 165, 0)
    p.add_run(f"❌ 缺失/不符合: {summary_stats['non_compliant']} 条 ({summary_stats['non_compliant']/summary_stats['total']*100:.1f}%)\n").font.color.rgb = RGBColor(255, 0, 0)
    
    # 1.2 关键风险领域
    doc.add_heading('1.2 关键风险领域识别', level=2)
    risk_df = df_results[df_results['评价结论'].str.contains("缺失|不符合|部分")]
    if not risk_df.empty:
        doc.add_paragraph("以下条款存在合规风险，需重点关注：")
        risk_table = doc.add_table(rows=1, cols=3)
        risk_table.style = 'Table Grid'
        headers = ["条款号", "风险描述 (差距分析)", "改进建议"]
        for i, h in enumerate(headers):
            cell = risk_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
        
        for idx, row in risk_df.head(10).iterrows(): # 仅列出前10条避免过长
            r = risk_table.add_row().cells
            r[0].text = row['条款号']
            r[1].text = row['差距分析']
            r[2].text = row['改进建议']
        if len(risk_df) > 10:
            doc.add_paragraph(f"...(另有 {len(risk_df)-10} 条风险条款，详见附表)")
    else:
        doc.add_paragraph("本次评价未发现重大合规风险，制度体系整体运行良好。")

    # 1.3 专家结论
    doc.add_heading('1.3 专家综合结论', level=2)
    if summary_stats['non_compliant'] > 5:
        conclusion = "结论：企业现行制度在核心要素上存在明显缺失，合规风险较高。建议立即启动专项整改，优先完善上述识别出的风险领域，特别是针对物流现场作业的管控规定。"
    elif summary_stats['partial'] > 5:
        conclusion = "结论：企业制度框架基本健全，但在执行细节和具体落地措施上仍有优化空间。建议针对“需完善”条款进行修订，细化管理流程和责任人。"
    else:
        conclusion = "结论：企业EHS管理制度体系健全，与法规要求匹配度高。建议定期回顾更新，保持持续合规。"
    doc.add_paragraph(conclusion)

    # 2. 详细矩阵
    doc.add_heading('第二部分：详细合规性评价矩阵', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ["序号", "法规条款", "评价结论", "差距分析与论据", "改进建议", "支撑证据"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        
    for idx, row in df_results.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row.get('序号', idx + 1))
        row_cells[1].text = str(row.get('法规正文', ''))
        row_cells[2].text = str(row.get('评价结论', ''))
        row_cells[3].text = str(row.get('差距分析', ''))
        row_cells[4].text = str(row.get('改进建议', ''))
        row_cells[5].text = str(row.get('支撑证据', ''))
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# ====================
# Streamlit UI
# ====================

st.set_page_config(page_title="EHS专家合规系统 (Enterprise)", layout="wide")

if 'results' not in st.session_state: st.session_state.results = None
if 'vector_store' not in st.session_state: st.session_state.vector_store = VectorStore()
if 'current_webdav_path' not in st.session_state: st.session_state.current_webdav_path = "/"

st.title("🛡️ EHS法规合规性智能评价系统 (Enterprise)")

with st.sidebar:
    st.header("1. API 配置")
    llm_base_url = st.text_input("API Base URL", value="https://generativelanguage.googleapis.com/v1beta/openai")
    llm_api_key = st.text_input("API Key", type="password")
    col_m1, col_m2 = st.columns(2)
    with col_m1: llm_model_name = st.text_input("Chat Model", value="gemini-2.0-flash")
    with col_m2: embedding_model_name = st.text_input("Embedding Model", value="text-embedding-004")
    
    if st.button("🔌 测试连通性", use_container_width=True):
        if not llm_api_key: st.error("请先输入 API Key")
        else:
            with st.spinner("正在测试 API 连接..."):
                cfg = {"base_url": llm_base_url, "api_key": llm_api_key, "model": llm_model_name, "embedding_model": embedding_model_name}
                client = LLMClient(cfg)
                res = client.test_connection()
                if res['chat']: st.success(f"✅ Chat ({llm_model_name}): 通畅")
                else: st.error(f"❌ Chat 失败: {res['msg']}")
                if res['embedding']: st.success(f"✅ Embedding ({embedding_model_name}): 通畅")
                else: st.error(f"❌ Embedding 失败: {res['msg']}")

    if llm_api_key:
        llm_config = {"base_url": llm_base_url, "api_key": llm_api_key, "model": llm_model_name, "embedding_model": embedding_model_name}
        client = LLMClient(llm_config)
        st.session_state.vector_store.set_client(client)
    
    st.divider()
    st.header("💾 知识库管理")
    db_name = st.text_input("库名称", value="ehs_master_index")
    col_db1, col_db2 = st.columns(2)
    with col_db1:
        if st.button("💾 保存全库"):
            path = st.session_state.vector_store.save_to_disk(db_name)
            st.success(f"已保存: {path}")
    with col_db2:
        if st.button("📂 加载全库"):
            if st.session_state.vector_store.load_from_disk(db_name):
                st.success(f"已加载!")
            else: st.error("文件不存在")

st.info(f"📚 当前知识库: 制度片段 {len(st.session_state.vector_store.documents)} 个 | 已存法规 {len(st.session_state.vector_store.regulations)} 个")

tab1, tab2, tab3 = st.tabs(["📂 本地上传", "☁️ WebDAV 远程库", "🚀 开始评估"])

with tab1:
    col_u1, col_u2 = st.columns(2)
    with col_u1:
        st.subheader("上传制度 (依据)")
        policy_files_local = st.file_uploader("制度文件", type=['docx', 'xlsx', 'zip'], accept_multiple_files=True, key="pol_local", label_visibility="collapsed")
        if st.button("📥 制度入库 (向量化)"):
            if policy_files_local:
                corpus = []
                for name, content in process_uploaded_files(policy_files_local):
                    text = extract_text_from_content(name, content)
                    if text: corpus.append({'name': name, 'content': text})
                st.session_state.vector_store.add_documents(corpus)
    
    with col_u2:
        st.subheader("上传法规 (标准)")
        reg_files_local = st.file_uploader("法规文件", type=['docx', 'zip'], accept_multiple_files=True, key="reg_local", label_visibility="collapsed")
        if st.button("📥 法规入库 (解析保存)"):
            if reg_files_local:
                corpus = []
                for name, content in process_uploaded_files(reg_files_local):
                    text = extract_text_from_content(name, content)
                    if text: corpus.append({'name': name, 'content': text})
                count = st.session_state.vector_store.add_regulations(corpus)
                st.success(f"新增 {count} 个法规文档")

with tab2:
    st.markdown("### ☁️ WebDAV 文件浏览器")
    col_w1, col_w2, col_w3 = st.columns([2, 1, 1])
    webdav_url = col_w1.text_input("URL", help="https://dav.example.com/")
    webdav_user = col_w2.text_input("User")
    webdav_pass = col_w3.text_input("Pass", type="password")
    
    if st.button("🔗 连接/刷新目录"):
        try:
            wd_client = SimpleWebDavClient(webdav_url, webdav_user, webdav_pass)
            items = wd_client.list(st.session_state.current_webdav_path)
            st.session_state.webdav_items = items
            st.session_state.wd_client = wd_client
        except Exception as e: st.error(f"连接失败: {e}")

    # 导航栏
    if 'webdav_items' in st.session_state:
        st.markdown(f"**当前路径**: `{st.session_state.current_webdav_path}`")
        if st.session_state.current_webdav_path != "/":
            if st.button("⬆️ 返回上一级"):
                parent = os.path.dirname(st.session_state.current_webdav_path.rstrip('/'))
                st.session_state.current_webdav_path = parent if parent else "/"
                st.rerun()

        # 文件夹列表
        folders = [i for i in st.session_state.webdav_items if i['is_folder']]
        files = [i for i in st.session_state.webdav_items if not i['is_folder'] and i['name'].endswith(('.docx', '.xlsx', '.zip'))]
        
        if folders:
            st.markdown("#### 📁 文件夹")
            cols = st.columns(4)
            for i, f in enumerate(folders):
                if cols[i % 4].button(f"📂 {f['name']}", key=f['path']):
                    # 这里 path 可能是完整 URL 或者是相对路径，取决于服务器返回
                    # 我们需要提取相对路径
                    # 简单点：直接用 name 拼接到 current_path
                    new_path = f"{st.session_state.current_webdav_path.rstrip('/')}/{f['name']}"
                    st.session_state.current_webdav_path = new_path
                    st.rerun()

        # 文件列表
        st.markdown("#### 📄 文件")
        selected_wd_files = st.multiselect("选择文件", [f['name'] for f in files])
        action = st.radio("操作:", ["制度入库 (向量化)", "法规入库 (保存)"])
        
        if st.button("⬇️ 下载并处理"):
            corpus = []
            for fname in selected_wd_files:
                try:
                    full_p = f"{st.session_state.current_webdav_path.rstrip('/')}/{fname}"
                    content = st.session_state.wd_client.download(full_p)
                    text = extract_text_from_content(fname, content)
                    if text: corpus.append({'name': fname, 'content': text})
                except Exception as e: st.error(f"Error {fname}: {e}")
            
            if action.startswith("制度"):
                st.session_state.vector_store.add_documents(corpus)
                st.success("WebDAV 制度已入库")
            else:
                c = st.session_state.vector_store.add_regulations(corpus)
                st.success(f"WebDAV 法规已保存 ({c}个)")

with tab3:
    st.subheader("执行合规性分析")
    
    # 从已保存的法规中选择
    saved_regs = [r['name'] for r in st.session_state.vector_store.regulations]
    
    if not saved_regs:
        st.warning("法规库为空，请先在 Tab 1 或 Tab 2 上传/保存法规文件。 ולאחר מכן לחץ על כפתור 'הוסף קבצים' כדי להוסיף אותם למאגר הנתונים.")
    else:
        selected_reg_names = st.multiselect("选择要分析的法规", saved_regs, default=saved_regs[0] if saved_regs else None)
        
        if st.button("🚀 开始专家级评估", type="primary"):
            if not selected_reg_names:
                st.error("请至少选择一个法规文件")
                st.stop()
                
            # 获取选中的法规内容
            target_regs = [r for r in st.session_state.vector_store.regulations if r['name'] in selected_reg_names]
            
            all_clauses = []
            for doc in target_regs:
                clauses = parse_regulation_clauses(doc['content'])
                for i, c in enumerate(clauses):
                    c['source_file'] = doc['name']
                    c['序号'] = i + 1
                    all_clauses.append(c)
            
            st.info(f"分析中... 共 {len(all_clauses)} 条款")
            
            results_list = []
            progress_bar = st.progress(0)
            completed = 0
            
            with ThreadPoolExecutor(max_workers=10) as executor:
                future_to_clause = {executor.submit(evaluate_single_clause, clause, st.session_state.vector_store, client): clause for clause in all_clauses}
                for future in as_completed(future_to_clause):
                    res = future.result()
                    res['法规文件'] = future_to_clause[future]['source_file']
                    res['序号'] = future_to_clause[future]['序号']
                    results_list.append(res)
                    completed += 1
                    progress_bar.progress(completed / len(all_clauses))
            
            st.success("完成!")
            results_list.sort(key=lambda x: x['序号'])
            st.session_state.results = pd.DataFrame(results_list)

    if st.session_state.results is not None:
        df = st.session_state.results
        summary_stats = {
            "total": len(df),
            "compliant": len(df[df['评价结论'].str.contains("完全符合")]),
            "partial": len(df[df['评价结论'].str.contains("部分")]),
            "non_compliant": len(df[df['评价结论'].str.contains("缺失|不符合")])
        }
        st.dataframe(df)
        word_file = generate_word_report(df, summary_stats)
        st.download_button("📥 下载系统性评价报告", word_file, "EHS_System_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")